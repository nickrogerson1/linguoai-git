import os
from io import BytesIO, StringIO
from html.parser import HTMLParser
from lxml import html, etree

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse

import pypandoc
from docx import Document
from copy import deepcopy
from bs4 import BeautifulSoup
from ..views import CorrectedSubmission, ImprovedSubmission, IeltsWritingTask2
from ..api_funcs.corrections import find_difference

from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.shared import RGBColor
from docx.shared import Pt

import pathlib

# Strip out any HTML tags in strings
class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs= True
        self.text = StringIO()
    def handle_data(self, d):
        self.text.write(d)
    def get_data(self):
        return self.text.getvalue()
    
def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()

def suffix(d):
    return {1:'st',2:'nd',3:'rd'}.get(d%20, 'th')

def custom_strftime(t):
    date_format = '%A {S} %B %Y at %-I:%M %p'
    return t.strftime(date_format).replace('{S}', str(t.day) + suffix(t.day))


@login_required(login_url="/login/")
def get_docx(request, pk, type=None, multi=False, sub=None):
    print(type)
    if type:
        type = int(type)

    sub_type = sub if sub else request.path.split('/')[1]
    user = request.user
    
    if sub_type == 'corrected-results':

        look_up = CorrectedSubmission.objects.get(pk=pk)
        sub = look_up.submission
        result = look_up.result
        date = custom_strftime(look_up.time_created)

        if(type):
            corrections = find_difference(sub, result)

            # Named capturing group attempt at doing this
            # pattern = r'(?P<type>(\s?=</span>)|(?<=<span class="diff_chg">\s)|(?<=<span class="diff_sub">\s)|(?<=<span class="diff_add">\s))'
            # # replacer = r'\g<type></span> |  <span class="diff_chg">|  <span class="diff_sub">| <span class="diff_add">'
            # replacer = r' \g<type>'

            # corrections = re.sub(pattern, replacer, corrections)

            corrections = (corrections.replace(' </span>', '</span> ')
                            .replace('<span class="diff_chg"> ', ' <span class="diff_chg">')
                            .replace('<span class="diff_sub"> ', ' <span class="diff_sub">')
                            .replace('<span class="diff_add"> ', ' <span class="diff_add">'))
            
            
        # Replace heading as colspan not supported
            table = html.fragment_fromstring(corrections)
            new_header = etree.fromstring('''
                <thead id="main-table">
                    <tr>
                        <th></th>
                        <th><strong>Original</strong></th>
                        <th></th>
                        <th><strong>Edited</strong></th>
                    </tr>
                </thead>
                ''')
            
            for el in table.getchildren():
                if el.tag == 'thead':
                    el.getparent().replace(el, new_header)

            modified_header = html.tostring(table).decode('utf-8')
            
        # Change the classes over to suit the needs of Pandoc
            changed_classes = (modified_header.replace('class="diff_chg"', 'custom-style="Red"')
                            .replace('class="diff_sub"', 'custom-style="Yellow"')
                            .replace('class="diff_add"', 'custom-style="Green"'))
            
        

            # print(changed_classes)

        # No other option but to save, reload and delete this file
            output = pypandoc.convert_text(source=changed_classes, format='html', to='markdown')
            filename = f'{user}-corrections-{pk}.docx'
        # Keep it in the copies directory although doesn't really matter
            file_path = f'{str(pathlib.Path(__file__).parent.resolve())}/{filename}'
            pypandoc.convert_text(source=output, format='markdown', to='docx', outputfile=file_path, extra_args=['--reference-doc=members/copies/docx_template.docx'])
            

        # Overwrite the existing doc to preserve the original styles
            doc = Document(file_path)
            doc_deep = deepcopy(doc)
      
        # Then get rid of it
            os.remove(file_path)
            
        # # Clear it then rebuild it
            for element in doc.element.body:
                element.clear()

        # Rebuild the doc
            doc.add_heading('Corrected Version\n').alignment = 1

    # This is a hack to mix low-level and high-level code which I don't really get
    # Otherwise the date will get shown before the table
            for element in doc_deep.element.body:
                doc.element.body.append(element)
                
                new_p = OxmlElement("w:p")
                element.addnext(new_p)
                new_para = Paragraph(new_p, element)
                run = new_para.add_run(f'\n\nThis submission was made on {date}')
                run.font.color.rgb = RGBColor.from_string('4F81BD')
                run.bold = True
                run.font.size = Pt(12)
                run.font.name = 'Calibri'
                break
                
            # Stream it and prepare for sending
            stream = BytesIO()
            doc.save(stream)
            stream.seek(0)

            if multi:
                return [stream, filename]

            response = HttpResponse(stream, content_type='application/vnd.openxmlformats')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            stream.close()
            return response
        
        else:
    # For the split Corrections
        # Build document
            doc = Document()
            doc.add_heading('Corrected Version')
            doc.add_paragraph(result)
            doc.add_page_break()
            doc.add_heading('Original Version')
            doc.add_paragraph(sub)
            doc.add_heading(f'This submission was made on {date}', 3)
            
        # Stream it and prepare for sending
            stream = BytesIO()
            doc.save(stream)
            stream.seek(0)
            
            filename = f'{user}-corrected-split-{pk}.docx'

            if multi:
                return [stream, filename]

            response = HttpResponse(stream, content_type='application/vnd.openxmlformats')
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
    
    elif sub_type == 'improved-results':

        look_up = ImprovedSubmission.objects.get(pk=pk)
        sub = strip_tags(look_up.submission)
        improved = strip_tags(look_up.improved_sub)
        date = custom_strftime(look_up.time_created)
    
    # Build document
        doc = Document()
        doc.add_heading('Improved Version')
        doc.add_paragraph(improved)
        doc.add_page_break()
        doc.add_heading('Original Version')
        doc.add_paragraph(sub)
        doc.add_heading(f'This submission was made on {date}', 3)
        
    # Stream it and prepare for sending
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        filename = f'{user}-improved-{pk}.docx'

        if multi:
                return [stream, filename]

        response = HttpResponse(stream, content_type='application/vnd.openxmlformats')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
    
    # Can only be IELTS left 
    else:
        
        look_up = IeltsWritingTask2.objects.get(pk=pk)
        question = strip_tags(look_up.question)
        answer = strip_tags(look_up.answer)
        
        band = look_up.band
        date = custom_strftime(look_up.time_created)

        soup = BeautifulSoup(look_up.score_res, 'html.parser')
        h3s = soup.find_all('h3')
        ps = soup.find_all('p')
        h5 = soup.find('h5')
    
    # Build document
        doc = Document()
        doc.add_heading(f'Ielts Writing Task 2 - Overall Band {band}', level=0)

        for i in range(4):
            doc.add_heading(h3s[i].string)
            doc.add_paragraph(ps[i].string)
        
        doc.add_heading(h5, level=3)
         
        doc.add_heading('Submitted Question')
        doc.add_paragraph(question)
        doc.add_heading('Your Original Answer')
        doc.add_paragraph(answer)
        doc.add_heading(f'This submission was made on {date} (UTC)', 3)
        
    # Stream it and prepare for sending
        stream = BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        filename = f'{user}-ielts-writing-task-2-{pk}-result.docx'

        if multi:
                return [stream, filename]

        response = HttpResponse(stream, content_type='application/vnd.openxmlformats')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response




from zipfile import ZipFile, ZIP_DEFLATED


def get_bulk_docx(request, pks, type=None):

    origin = request.path.split('/')[1]
    if origin == 'corrected-results':
        filename = f'{request.user}-pdf-corrections.zip'
    elif origin == 'improved-results':
        filename = f'{request.user}-pdf-improved.zip'
    else:
        filename = f'{request.user}-pdf-ielts-writing-task-2.zip'

    # Ignore last one as it's empty
    pks = pks.split('/')[:-1]
    print(f'PKS: {pks}')

    # if just one file, return as is 
    if len(pks) == 1:
        return get_docx(request, pks[0], type)

    # Otherwise zip them
    buffer = BytesIO()

    with ZipFile(buffer, 'w', ZIP_DEFLATED) as f:
        for pk in pks:  
            fetch_file = get_docx(request, pk, type, multi=True)
            doc = fetch_file[0]
            fn = fetch_file[1]
            f.writestr(fn, doc.getvalue())

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"',
        'Content-Type': 'application/zip'
    }
        
    return HttpResponse(buffer.getvalue(), headers=headers)





def get_bulk_mixed_docx(request, url_str):

    filename = f'{request.user}-docxs.zip'

    url_list = url_str.split("/")[:-1]

    # Split them into chunks
    pairs = [url_list[i:i+2] for i in range(0, len(url_list),2)]
    print(f'Pairs: {pairs}')

    # if just one file, return as is 
    if len(pairs) == 1:
        return get_docx(request, pairs[0][1], sub=pairs[0][0])

    # Otherwise zip them
    buffer = BytesIO()

    with ZipFile(buffer, 'w', ZIP_DEFLATED) as f:
        for pair in pairs:  
            fetch_file = get_docx(request, pair[1], type=True, multi=True, sub=pair[0])
            doc = fetch_file[0]
            fn = fetch_file[1]
            f.writestr(fn, doc.getvalue())

    headers = {
        'Content-Disposition': f'attachment; filename="{filename}"',
        'Content-Type': 'application/zip'
    }
        
    return HttpResponse(buffer.getvalue(), headers=headers)